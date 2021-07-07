from chinese_stroke_sorting import sort_by_stroke
from docx import Document
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION as WD_ORIENT
from docx.shared import Inches
from docx.shared import Cm, Pt


def sort_name_list(name_lists):
    nlist = []
    for stri in name_lists:
        strlist = stri.split(" ")
        for name in strlist:
            if name != "":
                nlist.append(name)
    # print(nlist)

    # 将单字重新组合为两字姓名
    try:
        new_name_list = []
        for a in range(len(nlist)):
            if len(nlist[0]) == 1:
                name = nlist[0] + nlist[1]
                new_name_list.append(name)
                nlist.pop(0)
                nlist.pop(0)
            else:
                name = nlist[0]
                new_name_list.append(name)
                nlist.pop(0)
    except:
        pass
    # print(new_name_list)
    return new_name_list

def split_name_list_per_12(sorted_list):
    # 按照12个为一组分开创建列表
    index = 0
    new_list = []
    new_write_list = []
    for name in sorted_list:
        new_list.append(name)
        index += 1
        if index == 12:
            new_write_list.append(new_list)
            new_list = []
            index = 0
        if name == sorted_list[-1]:
            new_write_list.append(new_list)
            new_list = []
            index = 0
    return new_write_list

def write_student_name_per12(sorted_list, output_docx):
    # 顺序写入段落并调整字号
    for para in sorted_list:
        paragraph = output_docx.add_paragraph()  # 每循环一次创建一个段落
        paragraph.paragraph_format.space_before = Pt(0)  # 段前间距
        paragraph.paragraph_format.space_after = Pt(0)  # 段后间距
        for name in para:
            if len(name) == 4:
                # print(name)
                run = paragraph.add_run(name)
                run.font.size = Pt(8)  # 字体大小设置，和word里面的字号相对应
            elif len(name) == 2:
                run = paragraph.add_run(name[0])
                run.font.size = Pt(10.5)  # 字体大小设置，和word里面的字号相对应
                run1 = paragraph.add_run("  ")
                run1.font.size = Pt(10.5)
                run2 = paragraph.add_run(name[1])
                run2.font.size = Pt(10.5)  # 字体大小设置，和word里面的字号相对应
            else:
                run = paragraph.add_run(name)
                run.font.size = Pt(10.5)  # 字体大小设置，和word里面的字号相对应
            run1 = paragraph.add_run("  ")
    return new_name_list


# 定义输入输出文件
input_docx = Document("sort_file.docx")
output_docx = Document()
output_docx.styles['Normal'].font.name = u'宋体'
output_docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
output_docx.save("output.docx")

# 按行读取文件并转化为列表
row_list = []
for row in input_docx.paragraphs:
    row = row.text
    row_list.append(row)
# print(row_list)

# 构建行标签字典
Adict = {}
for i in range(len(row_list)):
    # for para in row_list:
    if "毕业" in row_list[i] or "小学" in row_list[i]:
        Adict[str(i)] = "year_line"
    elif "班主任" in row_list[i] or "政治辅导员" in row_list[i]:
        Adict[str(i)] = "clss_line"
    elif "133中转入" in row_list[i]:
        Adict[str(i)] = "spec_line"
    elif row_list[i] == "" or row_list[i] == " ":
        Adict[str(i)] = "split_line"
    else:
        Adict[str(i)] = "name_line"
# print(Adict)
# exit()

# 提取数据并写入
name_list = []
for i in range(len(row_list)):
    # print(i)
    # paragraph = output_docx.add_paragraph()
    if Adict[str(i)] == "year_line":
        paragraph = output_docx.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)  # 段前间距
        paragraph.paragraph_format.space_after = Pt(0)  # 段后间距
        run = paragraph.add_run(row_list[i])
        run.font.size = Pt(14)  # 字体大小设置，和word里面的字号相对应
        run.bold = True  # 文字加粗
    elif Adict[str(i)] == "spec_line":
        paragraph = output_docx.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.space_before = Pt(0)  # 段前间距
        paragraph.paragraph_format.space_after = Pt(0)  # 段后间距
        run = paragraph.add_run(row_list[i])
        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
        run.bold = True  # 文字加粗
    elif Adict[str(i)] == "clss_line":
        paragraph = output_docx.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)  # 段前间距
        paragraph.paragraph_format.space_after = Pt(0)  # 段后间距
        run = paragraph.add_run(row_list[i])
        run.font.size = Pt(10.5)  # 字体大小设置，和word里面的字号相对应
        run.bold = True  # 文字加粗
        # print(row_list[i])
    elif Adict[str(i)] == "name_line":
        name_list.append(row_list[i])
    elif Adict[str(i)] == "split_line":
        new_name_list = split_name_list_per_12(sort_by_stroke(sort_name_list(name_list)))
        write_student_name_per12(new_name_list, output_docx)
        if new_name_list[-1] != []:
            paragraph_em = output_docx.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)  # 段前间距
            paragraph.paragraph_format.space_after = Pt(0)  # 段后间距
            paragraph_em.add_run("")
        print(new_name_list)
        name_list = []

    output_docx.save("output.docx")
